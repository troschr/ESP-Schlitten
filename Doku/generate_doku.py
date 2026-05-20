"""
Generiert die Softwaredokumentation für das Plattenwechsler-System als Word-Dokument.
Ausführen: python3 generate_doku.py
"""

import os
import glob
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Pfade ───────────────────────────────────────────────────────────────────
SCRIPT_DIR   = os.path.dirname(os.path.abspath(__file__))
ESP_SRC_DIR  = os.path.join(SCRIPT_DIR, "..", "src")
OUTPUT_FILE  = os.path.join(SCRIPT_DIR, "Softwaredokumentation_Plattenwechsler.docx")

# ─── Hilfsfunktionen ─────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.style.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p

def add_body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
    return p

def add_code_block(doc, text, caption=None):
    if caption:
        cp = doc.add_paragraph(caption)
        cp.paragraph_format.space_after = Pt(2)
        for run in cp.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F2F2F2")
        p._p.get_or_add_pPr().append(shading)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)

    doc.add_paragraph()

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    # Header
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].runs[0].font.name = "Calibri"
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D9E1F2")
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i+1].cells[c_i]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].runs[0].font.name = "Calibri"
    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def page_break(doc):
    doc.add_page_break()

def add_source_file(doc, filepath, label):
    """Fügt eine Quelldatei mit Header und Code-Block zum Anhang hinzu."""
    p = doc.add_heading(label, level=3)
    p.paragraph_format.space_before = Pt(12)
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            content = f.read()
    except Exception:
        content = "(Datei nicht gefunden)"
    # Ausgabe als Code-Block, aber ohne add_code_block (zu viele Absätze bei großen Dateien)
    # Stattdessen: pro Zeile ein Run in einem Absatz
    for line in content.split("\n"):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        p2.paragraph_format.left_indent  = Cm(0.3)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F7F7F7")
        p2._p.get_or_add_pPr().append(shading)
        run = p2.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(7.5)

# ─── Dokument aufbauen ───────────────────────────────────────────────────────

doc = Document()

# Seitenränder
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Titelseite ──────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Cm(4)
run = title_p.add_run("Softwaredokumentation")
set_font(run, size=26, bold=True, color=(0x1F, 0x39, 0x64))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run("Automatischer Plattenwechsler – ESP32-Firmware & Raspberry Pi Software")
set_font(run2, size=14, color=(0x40, 0x40, 0x40))

doc.add_paragraph()
doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = info_p.add_run("DHBW Stuttgart · Kurs G6-TWIE23A\nMai 2026")
set_font(run3, size=11, color=(0x60, 0x60, 0x60))

page_break(doc)

# ── Inhaltsverzeichnis (manuell) ────────────────────────────────────────────
add_heading(doc, "Inhaltsverzeichnis", level=1)
toc_items = [
    ("1", "Systemarchitektur", "3"),
    ("1.1", "Gesamtübersicht", "3"),
    ("1.2", "ESP32-Firmware", "3"),
    ("1.3", "Raspberry Pi Software", "4"),
    ("2", "Kommunikationsstruktur", "4"),
    ("2.1", "ESP ↔ Pi: UART-Protokoll", "4"),
    ("2.2", "Pi ↔ Außenwelt: MQTT, GPIO, Telegram", "6"),
    ("3", "Ablaufsteuerung", "6"),
    ("3.1", "Zustandsmaschine ESP32", "6"),
    ("3.2", "Zustandsmaschine Raspberry Pi (Hauptablauf)", "8"),
    ("3.3", "Plattenwechsel-Sequenz im Detail", "9"),
    ("Anhang A", "ESP32-Quellcode", "11"),
    ("Anhang B", "Raspberry Pi – Quelldateien (Übersicht)", "—"),
]
for num, title, page in toc_items:
    tp = doc.add_paragraph()
    tp.paragraph_format.space_after = Pt(3)
    run = tp.add_run(f"{num}   {title}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    if num in ("1", "2", "3"):
        run.font.bold = True

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 1. SYSTEMARCHITEKTUR
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "1  Systemarchitektur", level=1)

add_heading(doc, "1.1  Gesamtübersicht", level=2)
add_body(doc,
    "Das Plattenwechsler-System besteht aus zwei eigenständigen Softwarekomponenten, "
    "die über eine serielle UART-Verbindung (USB) miteinander kommunizieren: einer "
    "Echtzeit-Firmware auf einem ESP32-Mikrocontroller und einer Python-Anwendung auf "
    "einem Raspberry Pi. Beide Komponenten sind klar voneinander getrennt – der ESP32 "
    "übernimmt ausschließlich die hardwarenahe Motorsteuerung und Sensorauswertung, "
    "der Raspberry Pi die übergeordnete Auftragslogik, die Benutzeroberflächen und die "
    "Anbindung an externe Systeme."
)

add_body(doc,
    "Der Raspberry Pi ist über weitere Schnittstellen mit der Umgebung verbunden: GPIO-"
    "Eingänge empfangen Signale der angeschlossenen 3D-Drucker (Druckende-Meldung) und "
    "der mechanischen Endschalter der X- und Z-Achse. Ein MQTT-Broker ermöglicht die "
    "Integration in Home-Automation- und Fabriksteuerungs-Infrastrukturen. Ein Telegram-"
    "Bot erlaubt Fernbedienung über Mobilgerät. Lokal steht eine PyQt5-Vollbild-HMI sowie "
    "ein Flask-Webinterface zur Verfügung."
)

add_table(doc,
    headers=["Komponente", "Technologie", "Aufgabe"],
    rows=[
        ["ESP32-Firmware",    "C++, Arduino Framework, PlatformIO", "Motorsteuerung, Sensorik, Protokoll-Handler"],
        ["Pi-Hauptanwendung", "Python 3, PyQt5",                    "Auftragslogik, HMI, MQTT, Telegram, GPIO"],
        ["Pi-Webapp",         "Python 3, Flask, Server-Sent Events","Browser-basiertes Bedieninterface"],
        ["UART-Verbindung",   "115200 Baud, 8N1, ASCII",            "Befehl/Antwort-Protokoll Pi ↔ ESP"],
        ["MQTT-Broker",       "Paho MQTT, QoS 1",                   "Status-Verteilung, Remote-Kommandos"],
    ],
    col_widths=[4.5, 5.5, 7.0]
)

add_heading(doc, "1.2  ESP32-Firmware", level=2)
add_body(doc,
    "Die Firmware ist in C++ für das Arduino-Framework strukturiert und in mehrere "
    "Schichten aufgeteilt. Der Einstiegspunkt ist `src/main.cpp`, der lediglich eine "
    "Instanz des `AppController` erzeugt und dessen `begin()`- und `update()`-Methoden "
    "in `setup()` und `loop()` aufruft."
)

add_body(doc,
    "Die Schichtenstruktur der Firmware gliedert sich wie folgt:"
)
add_bullet(doc, "config/ – Statische Konfiguration: Pins.h (Pinbelegung) und Config.h (alle Betriebsparameter als constexpr)")
add_bullet(doc, "core/Types.h – Gemeinsame Typdefinitionen: AppState, ErrorCode, Command, Position, Snapshots")
add_bullet(doc, "drivers/ – Hardwaretreiber: ClMotor (CL42T Closed-Loop-Stepper) und DrvActuator (DRV8825)")
add_bullet(doc, "sensors/ – SensorManager: VL53L0X Türsensor (I2C) und TF-Luna LiDAR (I2C)")
add_bullet(doc, "comm/ – CommandInterface: UART-Eingabe, Protokoll-Parser, Kommando-Queue")
add_bullet(doc, "status/ – StatusReporter: formatierte Ausgabe aller RSP/EVT-Nachrichten")
add_bullet(doc, "app/AppController – Zentrale Steuerklasse: Zustandsmaschine, Kommando-Dispatcher, Bewegungsabläufe")

doc.add_paragraph()
add_body(doc,
    "Beide Motortypen werden non-blocking betrieben: ClMotor implementiert eine "
    "trapezförmige Geschwindigkeitsrampe mit Microsekunden-genauer Schrittzählung über "
    "`micros()`. DrvActuator steuert DRV8825-Treiber mit konstantem Schrittverzögerung. "
    "Alle Achsen werden in jedem Loop-Durchlauf durch Aufruf von `update()` fortgeschrieben "
    "– blockierende Delays existieren nicht."
)

add_heading(doc, "1.3  Raspberry Pi Software", level=2)
add_body(doc,
    "Die Pi-Anwendung ist als Python-Paket `plattenwechsler` organisiert. Der Einstiegspunkt "
    "`main.py` akzeptiert CLI-Flags (`--mock`, `--no-gui`, `--no-mqtt`, `--no-telegram`) und "
    "instanziiert alle Subsysteme. Die Kernkomponente ist `Hauptablauf` in "
    "`core/hauptablauf.py`, eine Zustandsmaschine, die alle Betriebsphasen orchestriert."
)

add_body(doc,
    "Alle I/O-Kanäle sind als separate Module gekapselt und werden dem Hauptablauf als "
    "Abhängigkeiten injiziert. Rückmeldungen erfolgen ausschließlich über Callbacks und "
    "threading.Event-Objekte. Das Designprinzip ist strikte Trennung: kein I/O-Modul kennt "
    "Auftragslogik, kein Auftragsmodul greift direkt auf Hardware zu."
)

add_table(doc,
    headers=["Modul", "Datei", "Funktion"],
    rows=[
        ["Hauptablauf",   "core/hauptablauf.py",    "Zustandsmaschine, Plattenwechsel-Sequenz"],
        ["Fehlerbehandlung", "core/fehler.py",       "Fehlerverwaltung, Quittierung, Historie"],
        ["Auftragswarteschlange", "core/auftrag_queue.py", "FIFO-Queue mit Deduplizierung"],
        ["ESP-Kommunikation", "io_/esp_client.py",   "UART-Protokoll, Heartbeat, Reconnect"],
        ["GPIO-Manager",  "io_/gpio_manager.py",     "gpiozero-Wrapper, Endschalter, E-Stop"],
        ["MQTT-Client",   "io_/mqtt_client.py",      "Broker-Verbindung, Status-Publish, Kommando-Subscribe"],
        ["Telegram-Bot",  "io_/telegram_client.py",  "Fernbedienung per Messenger"],
        ["PyQt5-HMI",     "ui/main_window.py",       "Vollbild-Touchscreen-Oberfläche"],
        ["Webapp",        "webapp/app.py",            "Flask-Webinterface mit SSE"],
    ],
    col_widths=[4.5, 5.0, 7.5]
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 2. KOMMUNIKATIONSSTRUKTUR
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "2  Kommunikationsstruktur", level=1)

add_heading(doc, "2.1  ESP ↔ Pi: UART-Protokoll", level=2)
add_body(doc,
    "Die gesamte Kommunikation zwischen Raspberry Pi und ESP32 läuft über eine "
    "UART-Verbindung (115200 Baud, 8N1, ASCII). Pro Zeile wird genau eine Nachricht "
    "übertragen, Felder sind durch Semikolon getrennt. Das Protokoll unterscheidet "
    "drei Nachrichtentypen:"
)

add_table(doc,
    headers=["Typ", "Richtung", "Format", "Bedeutung"],
    rows=[
        ["CMD", "Pi → ESP", "CMD;<id>;<verb>[;<key>=<val>…]", "Befehl vom Pi an den ESP"],
        ["RSP", "ESP → Pi", "RSP;<id>;ACK  oder  RSP;<id>;ERR;<code>", "Sofortantwort auf CMD"],
        ["EVT", "ESP → Pi", "EVT;<id>;<event>[;<key>=<val>…]", "Asynchrones Ereignis"],
    ],
    col_widths=[1.8, 2.2, 6.0, 4.0]
)

add_body(doc,
    "Die ID in CMD/RSP ist eine vom Pi frei vergebene, aufsteigende Ganzzahl ≥ 0. "
    "Der ESP spiegelt sie in seiner Antwort zurück, wodurch der Pi Antworten eindeutig "
    "ausstehenden Kommandos zuordnen kann. EVT-Nachrichten ohne direkten Befehlsbezug "
    "verwenden immer ID 0."
)

add_body(doc,
    "Der Pi sendet ein Kommando und wartet zunächst auf RSP;ACK (Befehl angenommen) "
    "oder RSP;ERR (abgelehnt). Für lang laufende Aktionen (Homing, Fahrt, Plattenwechsel) "
    "folgt nach dem ACK ein asynchrones EVT;OK;…_DONE sobald die Aktion abgeschlossen ist. "
    "Alle Befehle, die einen Bewegungsablauf starten, sind zustandsabhängig: der ESP "
    "lehnt sie mit ERR;BUSY oder ERR;NOT_REFERENCED ab, wenn die Vorbedingungen nicht erfüllt sind."
)

add_body(doc, "Die vollständige Befehlsliste:")

add_table(doc,
    headers=["Befehl", "Parameter", "Abschluss-Event"],
    rows=[
        ["PING",             "–",                                             "PONG"],
        ["STATUS",           "–",                                             "EVT;STATUS"],
        ["HOME",             "–",                                             "HOME_DONE"],
        ["MOVE_HOME",        "–",                                             "MOVE_HOME_DONE"],
        ["MOVE_TO",          "x=<mm>;z=<mm>",                                "MOVE_DONE"],
        ["PICKUP",           "gripper_depth=<mm>;lift_offset=<mm>",          "PICKUP_DONE"],
        ["DEPOSIT",          "gripper_depth=<mm>;lift_offset=<mm>",          "DEPOSIT_DONE"],
        ["OPEN_DOOR",        "x_approach;z_approach;arm_extend;radius;angle;hook_drop", "DOOR_OPEN_DONE"],
        ["CLOSE_DOOR",       "x_approach;z_approach;arm_extend;radius;angle;hook_drop", "DOOR_CLOSE_DONE"],
        ["HOME_SWITCH_HIT",  "axis=X|Z",                                     "– (intern)"],
        ["STOP",             "–",                                             "STOPPED"],
        ["RESET_ERROR",      "–",                                             "ERROR_RESET"],
        ["ASSUME_POSITION",  "–",                                             "ASSUME_POSITION_DONE"],
    ],
    col_widths=[4.0, 6.5, 5.5]
)

add_body(doc,
    "Besonderheit der Referenzfahrt: Die Endschalter der X- und Z-Achse sind am "
    "Raspberry Pi angeschlossen (GPIO-Eingänge). Der Pi überwacht diese Pins und "
    "sendet bei Auslösung sofort HOME_SWITCH_HIT;axis=X bzw. Z. Der ESP stoppt "
    "daraufhin den betreffenden Motor und setzt die Achsposition auf 0 mm. "
    "Greifer- und Türarm-Endschalter sind hingegen direkt am ESP angeschlossen "
    "und werden intern ausgewertet."
)

add_body(doc,
    "Der ESP sendet zusätzlich periodisch zwei spontane Nachrichten: einen HEARTBEAT "
    "alle 1000 ms (Lebenszeichen mit Uptime, Zustand und Position) sowie ein STATUS-"
    "Event bei jedem Zustandswechsel und im Stream-Modus alle 100 ms. Der Pi überwacht "
    "den Heartbeat und erkennt einen Verbindungsverlust nach konfigurierbarem Timeout."
)

add_body(doc,
    "Sonderfall MOVE_TO aus Home-Position (x=0, z=0): Bevor die eigentliche Fahrt "
    "beginnt, führt der ESP automatisch einen Z-Scan durch – die Z-Achse fährt die "
    "gesamte Verfahrlänge nach oben, während der TF-Luna-Hindernissensor kontinuierlich "
    "abgefragt wird. Erst nach fehlerfreiem Scan startet die Zielfahrt. Dieser Mechanismus "
    "verhindert Kollisionen mit offen stehenden Druckertüren oder anderem Zubehör im Fahrweg."
)

add_heading(doc, "2.2  Pi ↔ Außenwelt: MQTT, GPIO, Telegram", level=2)

add_body(doc,
    "Der Raspberry Pi stellt über MQTT-Topics eine standardisierte Schnittstelle für "
    "externe Systeme bereit. Status-Topics werden mit Retained-Flag publiziert, sodass "
    "neue Subscriber sofort den aktuellen Zustand erhalten. Kommandos werden über "
    "separierte Topics empfangen."
)

add_table(doc,
    headers=["Topic (Basis: plattenwechsler/)", "Richtung", "Inhalt"],
    rows=[
        ["status/online",          "Pi → Broker", "\"online\" / \"offline\" (retained)"],
        ["status/system",          "Pi → Broker", "Vollständiger System-Snapshot (JSON, retained)"],
        ["status/esp",             "Pi → Broker", "ESP-Statussnapshot (JSON, retained)"],
        ["status/drucker/<id>",    "Pi → Broker", "Pro-Drucker-Status (JSON, retained)"],
        ["error",                  "Pi → Broker", "Aktiver Fehler (JSON)"],
        ["event/auftrag",          "Pi → Broker", "Auftrags-Start/-Ende-Event (JSON)"],
        ["cmd/auftrag",            "Broker → Pi", "Plattenwechsel starten (Drucker-ID)"],
        ["cmd/referenzfahrt",      "Broker → Pi", "HOME-Befehl auslösen"],
        ["cmd/stop",               "Broker → Pi", "STOP-Befehl"],
        ["cmd/service/…",          "Broker → Pi", "Service-Modus-Steuerung"],
    ],
    col_widths=[6.0, 3.0, 7.0]
)

add_body(doc,
    "GPIO-Eingänge am Pi werden über die gpiozero-Bibliothek mit 50 ms Entprellung "
    "überwacht. Drei Eingangstypen sind konfigurierbar: Drucker-Fertig-Signale (je Drucker "
    "ein GPIO-Pin, active-low), Endschalter X und Z (active-low, liefern HOME_SWITCH_HIT), "
    "und ein Not-Aus-Eingang (konfigurierbarer Pin, invertierbar). Bei Not-Aus stoppt der "
    "Hauptablauf sofort und sendet STOP an den ESP."
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# 3. ABLAUFSTEUERUNG
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "3  Ablaufsteuerung", level=1)

add_heading(doc, "3.1  Zustandsmaschine ESP32", level=2)
add_body(doc,
    "Der AppController implementiert eine explizite Zustandsmaschine. Jeder Zustand "
    "entspricht einer Phase der Schlittensteuerung. Zustandsübergänge erfolgen ausschließlich "
    "über die interne Methode `setState()`, die gleichzeitig das EVT;STATE-Event an den Pi sendet."
)

add_table(doc,
    headers=["Zustand", "Bedeutung", "Übergang nach"],
    rows=[
        ["NOT_REFERENCED", "Gestartet, keine Referenzfahrt", "BUSY_HOMING (HOME), READY (ASSUME_POSITION)"],
        ["READY",          "Referenziert, wartet auf Befehl", "Alle BUSY_*-Zustände"],
        ["BUSY_HOMING",    "Referenzfahrt aller 4 Achsen",   "READY (HOME_DONE) / ERROR"],
        ["BUSY_SCANNING",  "Z-Scan vor Fahrt aus Home",       "BUSY_MOVING (scan ok) / ERROR"],
        ["BUSY_MOVING",    "Zielpositionsfahrt X+Z",          "READY (MOVE_DONE) / ERROR"],
        ["BUSY_MOVE_HOME", "Referenzrückfahrt X+Z",           "READY (MOVE_HOME_DONE) / ERROR"],
        ["BUSY_PICKUP",    "Plattenentnahme (3 Phasen)",      "READY (PICKUP_DONE) / ERROR"],
        ["BUSY_DEPOSIT",   "Plattenablage (4 Phasen)",        "READY (DEPOSIT_DONE) / ERROR"],
        ["BUSY_OPEN_DOOR", "Türöffnung (7 Phasen)",           "READY (DOOR_OPEN_DONE) / ERROR"],
        ["BUSY_CLOSE_DOOR","Türschließung (7 Phasen)",        "READY (DOOR_CLOSE_DONE) / ERROR"],
        ["STOPPED",        "Motor gestoppt via STOP",         "READY, BUSY_HOMING, BUSY_MOVING"],
        ["ERROR",          "Fehler, alle Motoren gestoppt",   "NOT_REFERENCED (RESET_ERROR)"],
    ],
    col_widths=[4.2, 5.8, 6.0]
)

add_body(doc,
    "Kommando-Dispatching: `CommandInterface::poll()` liest eingehende Bytes und "
    "puffert diese bis zum Zeilenende. Vollständige Zeilen werden durch `parseLine()` "
    "in `Command`-Objekte umgewandelt und in eine Queue mit 8 Einträgen eingereiht. "
    "`processCommands()` leert die Queue und ruft pro Befehl den zuständigen Handler auf. "
    "Alle Handler prüfen zuerst den aktuellen Zustand und senden bei Verletzung sofort "
    "RSP;ERR zurück – kein Motor wird gestartet."
)

add_body(doc,
    "Positionslimits: Vor jedem Bewegungsbefehl (MOVE_TO, OPEN_DOOR, CLOSE_DOOR, PICKUP, "
    "DEPOSIT) prüft der AppController, ob alle Zielpositionen innerhalb der konfigurierten "
    "Verfahrgrenzen liegen (0 mm bis MAX_TRAVEL_MM für X und Z, keine negativen Werte). "
    "Verletzungen werden mit RSP;ERR;POSITION_ERROR quittiert, ohne dass ein Motor startet."
)

add_body(doc,
    "Kreisbogen-Steuerung (Türöffnung/-schließung): Der Türöffnungsablauf koordiniert "
    "den Türarm (DRV8825) und die X-Achse (CL42T) in einem 7-phasigen Ablauf. Der "
    "Kreisbogen in Phase 3 wird durch Waypoints parametriert: Der Arm treibt das Timing "
    "mit konstantem Schritt-Delay, die X-Achse erhält bei jedem abgeschlossenen Arm-Schritt "
    "ein aktualisiertes Ziel. Da ClMotor Retargeting bei gleicher Fahrtrichtung unterstützt, "
    "hält X kontinuierlich Fahrt ohne Stop-Start zwischen Waypoints."
)

add_heading(doc, "3.2  Zustandsmaschine Raspberry Pi (Hauptablauf)", level=2)
add_body(doc,
    "Der Hauptablauf (`Hauptablauf`-Klasse in `core/hauptablauf.py`) verwaltet den "
    "Gesamtsystemzustand mit sechs Systemzuständen. Die interne Worker-Schleife läuft "
    "in einem Hintergrund-Thread und blockiert auf Events, wenn kein Auftrag anliegt. "
    "Alle Hardware-Callbacks (GPIO, ESP-Events) injizieren Events in diese Schleife."
)

add_table(doc,
    headers=["Systemzustand", "Beschreibung"],
    rows=[
        ["INIT",           "Systemstart, Initialisierung der Subsysteme"],
        ["REFERENZFAHRT",  "Wartet auf HOME_DONE vom ESP, Pi sendet HOME und überwacht Endschalter"],
        ["BEREITSCHAFT",   "System bereit, verarbeitet Aufträge aus der Queue"],
        ["PLATTENWECHSEL", "Aktiver Plattenwechsel-Ablauf (5 Hauptphasen)"],
        ["SERVICE",        "Manuelle Positionierung, kein Auftragsprocessing"],
        ["FEHLER",         "Fehlerbehandlung, wartet auf Benutzer-Quittierung"],
        ["NOT_AUS",        "Not-Aus ausgelöst, alle Aktionen gestoppt"],
    ],
    col_widths=[4.5, 11.5]
)

add_body(doc,
    "Fehlerbehandlung: Alle ESP-Fehler und Kommunikationsfehler landen im zentralen "
    "`FehlerBehandlung`-Objekt. Dort werden doppelte Fehler innerhalb von 2 Sekunden "
    "unterdrückt, eine Fehlerhistorie der letzten 100 Einträge wird gepflegt. Nach einem "
    "Fehler zeigt die HMI einen Dialog mit Optionen: Referenzfahrt neu starten (ja/nein) "
    "und ausstehende Aufträge löschen (ja/nein). Je nach Auswahl geht das System in "
    "REFERENZFAHRT oder direkt zurück in BEREITSCHAFT."
)

add_body(doc,
    "Auftragswarteschlange: `AuftragQueue` ist eine threadsichere FIFO-Queue mit "
    "Pro-Drucker-Deduplizierung. Jeder neue Auftrag für einen Drucker, der bereits in "
    "der Queue steht, wird verworfen. Die Queue erlaubt Vorrangeinfügen (vorne) für "
    "Fehler-Recovery-Aufträge. Auftragsquellen sind GPIO (Drucker-Fertig-Signal), "
    "MQTT (Fernsteuerung), Telegram-Bot und die lokale HMI."
)

add_heading(doc, "3.3  Plattenwechsel-Sequenz im Detail", level=2)
add_body(doc,
    "Ein vollständiger Plattenwechsel gliedert sich in fünf Hauptphasen, die der Hauptablauf "
    "sequenziell abarbeitet. Für jede Phase sendet der Pi mindestens einen Befehl an den ESP "
    "und wartet auf das zugehörige Abschluss-Event, bevor die nächste Phase beginnt."
)

add_body(doc, "Phase 1 – Alte Platte entnehmen:")
add_bullet(doc, "MOVE_TO zur Position vor dem Drucker (X/Z aus DruckerConfig)")
add_bullet(doc, "OPEN_DOOR mit druckerspezifischen Kreisbogen-Parametern")
add_bullet(doc, "PICKUP mit gripper_depth und lift_offset aus DruckerConfig")
add_bullet(doc, "Alte Platte liegt jetzt auf dem Greifer")

doc.add_paragraph()
add_body(doc, "Phase 2 – Alte Platte in Ablage deponieren:")
add_bullet(doc, "MOVE_TO zur konfigurierten Ablage-Position (AblageConfig)")
add_bullet(doc, "DEPOSIT mit gripper_depth und lift_offset aus AblageConfig")
add_bullet(doc, "Ablage wird als belegt markiert")

doc.add_paragraph()
add_body(doc, "Phase 3 – Neue Platte aus Magazin holen:")
add_bullet(doc, "MOVE_TO zur Magazin-Position (MagazinConfig)")
add_bullet(doc, "PICKUP mit Magazin-Parametern")
add_bullet(doc, "Magazinplatz wird als verfügbar markiert")

doc.add_paragraph()
add_body(doc, "Phase 4 – Neue Platte in Drucker einlegen:")
add_bullet(doc, "MOVE_TO zurück zur Druckerposition")
add_bullet(doc, "DEPOSIT in den Drucker")

doc.add_paragraph()
add_body(doc, "Phase 5 – Tür schließen und aufräumen:")
add_bullet(doc, "CLOSE_DOOR mit denselben Parametern wie OPEN_DOOR")
add_bullet(doc, "Falls keine weiteren Aufträge in der Queue: MOVE_HOME")
add_bullet(doc, "Zustand → BEREITSCHAFT")

doc.add_paragraph()
add_body(doc,
    "Fehlerbehandlung während des Wechsels: Tritt während einer Phase ein ESP-Fehler auf "
    "(z.B. OBSTACLE, POSITION_ERROR), sendet der ESP EVT;ERR und wechselt in ERROR. Der "
    "Pi empfängt das Fehler-Event über den ESP-Client-Callback und wechselt in den FEHLER-"
    "Zustand. Nach Benutzerquittierung und ggf. Referenzfahrt kann der Auftrag neu in die "
    "Queue eingereiht werden."
)

add_body(doc,
    "Sonderfall Hinderniserkennung beim Z-Scan: Erkennt der TF-Luna-Sensor beim Z-Scan "
    "ein Hindernis, wechselt der ESP in ERROR mit Fehlercode OBSTACLE. Der Pi führt "
    "RESET_ERROR durch (ESP → NOT_REFERENCED) und kann nach Behebung des Hindernisses "
    "mit ASSUME_POSITION die letzte bekannte Schlittenposition übernehmen, ohne eine "
    "vollständige Referenzfahrt durchzuführen. Anschließend kann der MOVE_TO-Befehl "
    "wiederholt werden."
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# ANHANG A – ESP32 Quellcode
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Anhang A  –  ESP32-Quellcode", level=1)
add_body(doc,
    "Im Folgenden ist der vollständige Quellcode der ESP32-Firmware abgedruckt. "
    "Die Dateien sind nach Schichten geordnet: Konfiguration, Typen, Treiber, Sensoren, "
    "Kommunikation, Statusreporting und Hauptsteuerung."
)

ESP_FILES = [
    ("src/config/Pins.h",            "Pinbelegung (Pins.h)"),
    ("src/config/Config.h",          "Betriebsparameter (Config.h)"),
    ("src/core/Types.h",             "Typdefinitionen (Types.h)"),
    ("src/drivers/ClMotor.h",        "CL42T-Treiber Header (ClMotor.h)"),
    ("src/drivers/ClMotor.cpp",      "CL42T-Treiber Implementierung (ClMotor.cpp)"),
    ("src/drivers/DrvActuator.h",    "DRV8825-Treiber Header (DrvActuator.h)"),
    ("src/drivers/DrvActuator.cpp",  "DRV8825-Treiber Implementierung (DrvActuator.cpp)"),
    ("src/sensors/SensorManager.h",  "Sensor-Manager Header (SensorManager.h)"),
    ("src/sensors/SensorManager.cpp","Sensor-Manager Implementierung (SensorManager.cpp)"),
    ("src/comm/CommandInterface.h",  "Kommando-Schnittstelle Header (CommandInterface.h)"),
    ("src/comm/CommandInterface.cpp","Kommando-Schnittstelle Implementierung (CommandInterface.cpp)"),
    ("src/status/StatusReporter.h",  "Status-Reporter Header (StatusReporter.h)"),
    ("src/status/StatusReporter.cpp","Status-Reporter Implementierung (StatusReporter.cpp)"),
    ("src/app/AppController.h",      "AppController Header (AppController.h)"),
    ("src/app/AppController.cpp",    "AppController Implementierung (AppController.cpp)"),
    ("src/main.cpp",                 "Einstiegspunkt (main.cpp)"),
]

BASE = os.path.join(SCRIPT_DIR, "..")
for rel_path, label in ESP_FILES:
    full_path = os.path.join(BASE, rel_path)
    if os.path.exists(full_path):
        add_source_file(doc, full_path, f"A.{ESP_FILES.index((rel_path, label))+1}  {label}")
    else:
        p = doc.add_paragraph(f"[nicht gefunden: {rel_path}]")
        p.paragraph_format.space_after = Pt(6)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════════════
# ANHANG B – Pi Quelldatei-Übersicht
# ═══════════════════════════════════════════════════════════════════════════
add_heading(doc, "Anhang B  –  Raspberry Pi – Quelldateien (Übersicht)", level=1)
add_body(doc,
    "Der vollständige Python-Quellcode des Raspberry Pi ist im öffentlichen GitHub-"
    "Repository verfügbar: https://github.com/NubaGames/plattenwechsler"
)
add_body(doc, "Verzeichnisstruktur:")

PI_TREE = """plattenwechsler/
├── main.py                   Einstiegspunkt, CLI-Argumente, Initialisierung
├── config.py                 YAML-Konfigurationsverwaltung (threadsicher)
├── types.py                  Enums, Dataclasses, Fehlertypen
├── logger.py                 Logging (Datei + Konsole, Rotation)
├── core/
│   ├── hauptablauf.py        Hauptzustandsmaschine (~800 Zeilen)
│   ├── fehler.py             Fehlerverwaltung, -historie, Quittierung
│   └── auftrag_queue.py      FIFO-Warteschlange mit Deduplizierung
├── io_/
│   ├── esp_client.py         UART-Protokoll-Client (~600 Zeilen)
│   ├── mock_esp_client.py    Software-Mock für Tests
│   ├── gpio_manager.py       gpiozero-Wrapper: Endschalter, Drucker-Signale
│   ├── mqtt_client.py        MQTT-Anbindung (Paho)
│   └── telegram_client.py    Telegram-Bot-Integration
└── ui/
    └── main_window.py        PyQt5-Vollbild-HMI (6 Tabs, Touchscreen)

webapp/
├── app.py                    Flask-Webserver mit Server-Sent Events
├── templates/index.html      Responsive Web-UI
└── static/
    ├── app.js                Realtime-Updates via SSE
    └── style.css             Dark-Theme, DHBW-Akzentfarbe

tests/
├── test_protokoll.py         Protokoll-Parser-Tests
└── test_hauptablauf_mock.py  Integrationstests mit MockEspClient"""

add_code_block(doc, PI_TREE)

# ═══════════════════════════════════════════════════════════════════════════
# ANHANG B – Pi Quellcode (heruntergeladen von GitHub)
# ═══════════════════════════════════════════════════════════════════════════
import urllib.request

GITHUB_RAW = "https://raw.githubusercontent.com/NubaGames/plattenwechsler/main/"

PI_FILES = [
    ("plattenwechsler/main.py",              "B.1  Einstiegspunkt (main.py)"),
    ("plattenwechsler/config.py",            "B.2  Konfiguration (config.py)"),
    ("plattenwechsler/types.py",             "B.3  Typdefinitionen (types.py)"),
    ("plattenwechsler/logger.py",            "B.4  Logging (logger.py)"),
    ("plattenwechsler/core/hauptablauf.py",  "B.5  Hauptablauf-Statemachine (hauptablauf.py)"),
    ("plattenwechsler/core/fehler.py",       "B.6  Fehlerbehandlung (fehler.py)"),
    ("plattenwechsler/core/auftrag_queue.py","B.7  Auftrags-Queue (auftrag_queue.py)"),
    ("plattenwechsler/io_/esp_client.py",    "B.8  ESP-Client / Basisklasse (esp_client.py)"),
    ("plattenwechsler/io_/mock_esp_client.py","B.9  Mock-ESP (mock_esp_client.py)"),
    ("plattenwechsler/io_/gpio_manager.py",  "B.10 GPIO-Manager (gpio_manager.py)"),
    ("plattenwechsler/io_/mqtt_client.py",   "B.11 MQTT-Client (mqtt_client.py)"),
    ("plattenwechsler/io_/telegram_client.py","B.12 Telegram-Client (telegram_client.py)"),
    ("plattenwechsler/ui/main_window.py",    "B.13 PyQt5-HMI (main_window.py)"),
    ("webapp/app.py",                        "B.14 Flask-Webapp (webapp/app.py)"),
]

def fetch_pi_file(rel_path):
    url = GITHUB_RAW + rel_path
    try:
        with urllib.request.urlopen(url, timeout=10) as resp:
            return resp.read().decode("utf-8")
    except Exception as e:
        return f"# (Fehler beim Laden: {e})"

print("Lade Pi-Quellcode von GitHub …")
for rel_path, label in PI_FILES:
    print(f"  {rel_path}")
    content = fetch_pi_file(rel_path)
    p = doc.add_heading(label, level=3)
    p.paragraph_format.space_before = Pt(12)
    for line in content.split("\n"):
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(0)
        p2.paragraph_format.left_indent  = Cm(0.3)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F7F7F7")
        p2._p.get_or_add_pPr().append(shading)
        run = p2.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(7.5)

# ─── Speichern ───────────────────────────────────────────────────────────────
doc.save(OUTPUT_FILE)
print(f"Dokument gespeichert: {OUTPUT_FILE}")
